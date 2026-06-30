from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("D:/Pruebas_PAE")
OUTPUT_DIR = Path("C:/Users/user/OneDrive/Documentos/Pruebas")
OUTPUT_FILE = OUTPUT_DIR / "Reporte_Cobertura_Condiciones.docx"

SOURCE_FILE = ROOT / "backend/services/exam-service/src/services/liveAnswerRules.js"
TEST_FILE = ROOT / "backend/services/exam-service/src/tests/liveAnswerRules.conditionCoverage.test.js"
LOG_FILE = OUTPUT_DIR / "jest-liveAnswerRules-output.txt"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)

    for margin_name, value in {
        "top": top,
        "start": start,
        "bottom": bottom,
        "end": end,
    }.items():
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def format_table(table, header_fill="F2F4F7"):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"
    set_repeat_table_header(table.rows[0])
    for idx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.1
                for run in paragraph.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(9)
            if idx == 0:
                set_cell_shading(cell, header_fill)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True


def set_column_widths(table, widths):
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)


def add_code_block(doc, code):
    for line in code.splitlines():
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        run = paragraph.add_run(line if line else " ")
        run.font.name = "Consolas"
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(31, 31, 31)
    doc.add_paragraph()


def add_kv_table(doc, rows):
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Campo"
    table.rows[0].cells[1].text = "Detalle"
    for key, value in rows:
        cells = table.add_row().cells
        cells[0].text = key
        cells[1].text = value
    set_column_widths(table, [1.7, 4.8])
    format_table(table)
    doc.add_paragraph()


def add_case_table(doc):
    cases = [
        [
            "test_c1_true_c2_true_detecta_respuesta_repetida",
            "answers=[{userId:'42', questionId:'900'}], userId='42', currentQuestionId='900'",
            "C1=true, C2=true; la decision C1 && C2 devuelve true.",
            "expect(result).toBe(true); caso aprobado en Jest.",
        ],
        [
            "test_c1_true_c2_false_no_bloquea_otra_pregunta",
            "answers=[{userId:'42', questionId:'901'}], userId='42', currentQuestionId='900'",
            "C1=true, C2=false; la decision C1 && C2 devuelve false.",
            "expect(result).toBe(false); caso aprobado en Jest.",
        ],
        [
            "test_c1_false_c2_true_no_bloquea_otro_usuario",
            "answers=[{userId:'99', questionId:'900'}], userId='42', currentQuestionId='900'",
            "C1=false, C2=true; por cortocircuito la decision devuelve false.",
            "expect(result).toBe(false); caso aprobado en Jest.",
        ],
    ]
    table = doc.add_table(rows=1, cols=4)
    headers = ["Caso", "Datos de entrada", "Flujo logico", "Resultado real"]
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
    for row in cases:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
    set_column_widths(table, [1.65, 1.95, 1.45, 1.45])
    format_table(table)
    doc.add_paragraph()


def add_truth_table(doc):
    rows = [
        ["EJECUTADA", "V", "V", "true", "Caso 1: respuesta repetida"],
        ["EJECUTADA", "V", "F", "false", "Caso 2: misma persona, otra pregunta"],
        ["EJECUTADA", "F", "V", "false", "Caso 3: otra persona, misma pregunta"],
        ["OMITIDA", "F", "F", "false", "No necesaria para cobertura de condiciones"],
    ]
    table = doc.add_table(rows=1, cols=5)
    headers = ["Estado", "C1", "C2", "Decision C1 && C2", "Observacion"]
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
        if row[0] == "EJECUTADA":
            set_cell_shading(cells[0], "E2F0D9")
    set_column_widths(table, [1.15, 0.55, 0.55, 1.35, 2.9])
    format_table(table)
    note = doc.add_paragraph()
    note.add_run(
        "Las filas no marcadas se omiten intencionalmente: la cobertura de condiciones busca "
        "evaluar cada condicion atomica como verdadera y falsa, sin ejecutar una tabla exhaustiva "
        "cuando no aporta nueva evidencia al objetivo formal."
    )
    doc.add_paragraph()


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for style_name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def build_report():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    source_code = SOURCE_FILE.read_text(encoding="utf-8")
    test_code = TEST_FILE.read_text(encoding="utf-8")
    terminal_log = LOG_FILE.read_text(encoding="utf-8")

    doc = Document()
    configure_document(doc)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("Reporte de Cobertura de Condiciones")
    run.font.name = "Calibri"
    run.font.size = Pt(20)
    run.bold = True
    run.font.color.rgb = RGBColor(11, 37, 69)

    subtitle = doc.add_paragraph()
    subtitle.add_run("Proyecto PAE - Modulo Learning / Motor de partida en vivo").bold = True
    subtitle.add_run("\nFecha de ejecucion: 16/06/2026")

    doc.add_heading("1. Codigo base evaluado", level=1)
    doc.add_paragraph(
        "La decision compuesta validada corresponde a la regla que evita que un participante "
        "responda dos veces la misma pregunta activa. Esta regla fue centralizada para que los "
        "flujos de challenge y trivia consuman la misma logica."
    )
    add_code_block(doc, source_code)

    doc.add_heading("2. Analisis de condiciones atomicas", level=1)
    add_kv_table(
        doc,
        [
            ("Decision compuesta", "String(answer.userId) === normalizedUserId && String(answer.questionId) === normalizedQuestionId"),
            ("C1", "String(answer.userId) === normalizedUserId"),
            ("C2", "String(answer.questionId) === normalizedQuestionId"),
            ("Objetivo", "Evaluar C1 como V/F y C2 como V/F al menos una vez en ejecucion real."),
            ("Resultado esperado antes de ejecutar", "3 tests aprobados: uno positivo y dos negativos estrategicos."),
        ],
    )

    doc.add_heading("3. Suite automatizada integrada", level=1)
    doc.add_paragraph(f"Ruta del test: {TEST_FILE}")
    add_code_block(doc, test_code)

    doc.add_heading("4. Flujo por caso ejecutado", level=1)
    add_case_table(doc)

    doc.add_heading("5. Tabla de verdad estrategica", level=1)
    add_truth_table(doc)

    doc.add_heading("6. Evidencia de ejecucion real", level=1)
    doc.add_paragraph("Comando ejecutado desde D:\\Pruebas_PAE:")
    add_code_block(
        doc,
        "npx jest backend/services/exam-service/src/tests/liveAnswerRules.conditionCoverage.test.js --runInBand --verbose",
    )
    doc.add_paragraph("Codigo de salida observado: 0")
    doc.add_paragraph("Log real capturado:")
    add_code_block(doc, terminal_log)

    doc.add_heading("7. Conclusion de cobertura", level=1)
    doc.add_paragraph(
        "Tras la ejecucion real en el sistema, se ha alcanzado la cobertura al evaluar "
        "C1: usuario de la respuesta coincide con el usuario actual como (V, F) y "
        "C2: pregunta de la respuesta coincide con la pregunta activa como (V, F)."
    )
    doc.add_paragraph(
        "La tecnica formal valida exhaustivamente el comportamiento booleano atomico de la "
        "decision seleccionada. Los casos omitidos no reducen el rigor del objetivo de cobertura "
        "porque no agregan un nuevo valor booleano para C1 o C2."
    )

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("PAE - Evidencia automatizada de QA").font.size = Pt(9)

    doc.save(OUTPUT_FILE)
    return OUTPUT_FILE


if __name__ == "__main__":
    print(build_report())
