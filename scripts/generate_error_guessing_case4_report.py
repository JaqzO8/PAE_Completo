from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("D:/Pruebas_PAE")
WORKSPACE_DIR = ROOT / ".agents/workspace"
DOCS_DIR = Path("C:/Users/user/OneDrive/Documentos/Pruebas")
OUTPUT_NAME = "Reporte_Error_Guessing_Caso4.docx"
PRIMARY_OUTPUT = WORKSPACE_DIR / OUTPUT_NAME
COPY_OUTPUT = DOCS_DIR / OUTPUT_NAME

TEST_FILE = ROOT / "backend/services/exam-service/src/tests/liveAnswerRules.conditionCoverage.test.js"
SOURCE_FILE = ROOT / "backend/services/exam-service/src/services/liveAnswerRules.js"
LOG_FILE = WORKSPACE_DIR / "error-guessing-caso4-jest-output.txt"


CASE4_CODE = """    test('caso_4_error_guessing_respuesta_corrupta_null_no_debe_romper_la_regla', () => {
        expect(() => hasAlreadyAnswered({
            answers: [null],
            userId: '42',
            currentQuestionId: '900',
        })).not.toThrow();
    });"""


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_column_widths(table, widths):
    for row in table.rows:
        for index, width in enumerate(widths):
            row.cells[index].width = Inches(width)


def format_table(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.1
                for run in paragraph.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(9)
                    if row_index == 0:
                        run.bold = True
            if row_index == 0:
                set_cell_shading(cell, "F2F4F7")


def add_code_block(doc, text):
    for line in text.splitlines():
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        run = paragraph.add_run(line if line else " ")
        run.font.name = "Consolas"
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(31, 31, 31)
    doc.add_paragraph()


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for style_name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_summary_table(doc):
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Elemento"
    table.rows[0].cells[1].text = "Resultado observado"
    rows = [
        ("Dato inyectado", "answers: [null], userId: '42', currentQuestionId: '900'"),
        ("Resultado esperado", "La regla no debe lanzar excepcion y debe tratar la entrada corrupta como no coincidente."),
        ("Resultado real", "TypeError: Cannot read properties of null (reading 'userId')"),
        ("Contraste", "Casos 1, 2 y 3 pasan; Caso 4 falla y expone falta de validacion defensiva."),
    ]
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
    set_column_widths(table, [1.8, 4.7])
    format_table(table)
    doc.add_paragraph()


def build_report():
    WORKSPACE_DIR.mkdir(parents=True, exist_ok=True)
    DOCS_DIR.mkdir(parents=True, exist_ok=True)

    source_code = SOURCE_FILE.read_text(encoding="utf-8")
    full_test_code = TEST_FILE.read_text(encoding="utf-8")
    terminal_log = LOG_FILE.read_text(encoding="utf-8", errors="replace")

    doc = Document()
    configure_document(doc)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("Reporte Error Guessing - Caso 4")
    run.font.name = "Calibri"
    run.font.size = Pt(20)
    run.bold = True
    run.font.color.rgb = RGBColor(11, 37, 69)

    subtitle = doc.add_paragraph()
    subtitle.add_run("Proyecto PAE - Hallazgo exploratorio sobre cobertura de condiciones").bold = True
    subtitle.add_run("\nFecha de ejecucion: 16/06/2026")

    doc.add_heading("1. Diseno heuristico", level=1)
    doc.add_paragraph(
        "El Caso 4 aplica Error Guessing sobre un defecto comun en sistemas que persisten "
        "arreglos JSON: un elemento corrupto o nulo dentro de la coleccion de respuestas. "
        "La cobertura de condiciones no obligaba a probar este caso porque sus Casos 1, 2 y 3 "
        "solo forzaban valores V/F de C1 y C2 usando objetos bien formados."
    )
    add_summary_table(doc)

    doc.add_heading("2. Codigo base afectado", level=1)
    add_code_block(doc, source_code)

    doc.add_heading("3. Codigo exclusivo del Caso 4", level=1)
    add_code_block(doc, CASE4_CODE)

    doc.add_heading("4. Suite real integrada", level=1)
    doc.add_paragraph(f"Archivo de pruebas: {TEST_FILE}")
    add_code_block(doc, full_test_code)

    doc.add_heading("5. Evidencia empirica de ejecucion", level=1)
    doc.add_paragraph("Comando ejecutado desde D:\\Pruebas_PAE:")
    add_code_block(
        doc,
        "npx jest backend/services/exam-service/src/tests/liveAnswerRules.conditionCoverage.test.js --runInBand --verbose --no-color",
    )
    doc.add_paragraph("Codigo de salida observado: 1")
    doc.add_paragraph("Log real capturado:")
    add_code_block(doc, terminal_log)

    doc.add_heading("6. Veredicto tecnico", level=1)
    doc.add_paragraph(
        "Los Casos 1, 2 y 3 consiguieron el objetivo formal de cobertura de condiciones, "
        "pero construyeron una confianza limitada: asumieron que cada elemento de answers "
        "siempre tiene forma de objeto. El Caso 4 rompe esa suposicion con una entrada nula "
        "y demuestra que la regla puede lanzar una excepcion no controlada antes de evaluar "
        "la decision booleana. Este hallazgo no contradice la cobertura; la complementa al "
        "mostrar un riesgo de robustez que vive fuera de la matriz C1/C2."
    )

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("PAE - Error Guessing Caso 4").font.size = Pt(9)

    doc.save(PRIMARY_OUTPUT)
    doc.save(COPY_OUTPUT)
    return PRIMARY_OUTPUT, COPY_OUTPUT


if __name__ == "__main__":
    primary, copied = build_report()
    print(primary)
    print(copied)
