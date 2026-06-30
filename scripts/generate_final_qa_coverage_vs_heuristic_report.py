from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("D:/Pruebas_PAE")
DOCS_DIR = Path("C:/Users/user/OneDrive/Documentos/Pruebas")
WORKSPACE_DIR = ROOT / ".agents/workspace"
OUTPUT_NAME = "Reporte_Final_QA_Cobertura_vs_Heuristica.docx"
PRIMARY_OUTPUT = DOCS_DIR / OUTPUT_NAME
WORKSPACE_COPY = WORKSPACE_DIR / OUTPUT_NAME

SOURCE_FILE = ROOT / "backend/services/exam-service/src/services/liveAnswerRules.js"
TEST_FILE = ROOT / "backend/services/exam-service/src/tests/liveAnswerRules.conditionCoverage.test.js"
CONDITION_LOG_FILE = DOCS_DIR / "jest-liveAnswerRules-output.txt"
ERROR_GUESSING_LOG_FILE = WORKSPACE_DIR / "error-guessing-caso4-jest-output.txt"

FINAL_CONCLUSION = (
    "El contraste empírico demuestra que una cifra de 100% de Cobertura de Condiciones puede ser "
    "correcta y, al mismo tiempo, insuficiente como argumento de seguridad. Los Casos 1, 2 y 3 "
    "probaron con rigor matemático que la decisión C1 && C2 responde bien cuando las entradas "
    "tienen la forma esperada; sin embargo, el Caso 4 mostró que una entrada nula dentro de "
    "answers rompe la ejecución antes de que la lógica booleana pueda proteger el flujo. Por eso, "
    "la cobertura formal debe verse como una base de confianza, no como una garantía final: la "
    "intuición heurística añade presión sobre supuestos, datos corruptos y fallos humanos que las "
    "métricas estructurales no están obligadas a explorar."
)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_column_widths(table, widths):
    for row in table.rows:
        for index, width in enumerate(widths):
            row.cells[index].width = Inches(width)


def format_table(table, header_fill="F2F4F7"):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"
    set_repeat_table_header(table.rows[0])
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.1
                for run in paragraph.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(9)
                    if row_index == 0:
                        run.bold = True
            if row_index == 0:
                set_cell_shading(cell, header_fill)


def add_code_block(doc, text):
    for line in text.splitlines():
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        run = paragraph.add_run(line if line else " ")
        run.font.name = "Consolas"
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(31, 31, 31)
    doc.add_paragraph()


def clean_log(text):
    return (
        text.replace("â—", "*")
        .replace("â€º", ">")
        .replace("\x1b", "")
    )


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for style_name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_title(doc):
    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("Reporte Final QA: Cobertura Formal vs Heuristica")
    run.font.name = "Calibri"
    run.font.size = Pt(20)
    run.bold = True
    run.font.color.rgb = RGBColor(11, 37, 69)

    subtitle = doc.add_paragraph()
    subtitle.add_run("Proyecto PAE - Evidencia consolidada de pruebas").bold = True
    subtitle.add_run("\nFecha de elaboracion: 16/06/2026")


def add_truth_table(doc):
    table = doc.add_table(rows=1, cols=5)
    headers = ["Caso", "C1 usuario coincide", "C2 pregunta coincide", "Resultado C1 && C2", "Estado"]
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header

    rows = [
        ("Caso 1", "V", "V", "true", "EJECUTADO - detecta respuesta repetida"),
        ("Caso 2", "V", "F", "false", "EJECUTADO - misma persona, otra pregunta"),
        ("Caso 3", "F", "V", "false", "EJECUTADO - otra persona, misma pregunta"),
        ("No ejecutado", "F", "F", "false", "Omitido: no aporta nuevo valor atomico a C1/C2"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = value
        if row[0].startswith("Caso"):
            set_cell_shading(cells[4], "E2F0D9")

    set_column_widths(table, [0.9, 1.35, 1.35, 1.25, 1.65])
    format_table(table)
    doc.add_paragraph(
        "Los Casos 1, 2 y 3 fueron suficientes para que cada condicion atomica se evaluara "
        "como verdadera y falsa al menos una vez. La fila F/F se omite intencionalmente porque "
        "no agrega nueva evidencia al objetivo formal de cobertura de condiciones."
    )


def add_formal_results_table(doc):
    table = doc.add_table(rows=1, cols=4)
    headers = ["Caso", "Entrada clave", "Flujo validado", "Resultado empirico"]
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header

    rows = [
        ("Caso 1", "answers=[{userId:'42', questionId:'900'}]", "C1=V, C2=V", "Aprobado: devuelve true"),
        ("Caso 2", "answers=[{userId:'42', questionId:'901'}]", "C1=V, C2=F", "Aprobado: devuelve false"),
        ("Caso 3", "answers=[{userId:'99', questionId:'900'}]", "C1=F, C2=V", "Aprobado: devuelve false"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = value

    set_column_widths(table, [0.9, 2.15, 1.3, 2.15])
    format_table(table)


def build_report():
    DOCS_DIR.mkdir(parents=True, exist_ok=True)
    WORKSPACE_DIR.mkdir(parents=True, exist_ok=True)

    source_code = SOURCE_FILE.read_text(encoding="utf-8")
    test_code = TEST_FILE.read_text(encoding="utf-8")
    condition_log = clean_log(CONDITION_LOG_FILE.read_text(encoding="utf-8", errors="replace"))
    error_log = clean_log(ERROR_GUESSING_LOG_FILE.read_text(encoding="utf-8", errors="replace"))

    doc = Document()
    configure_document(doc)
    add_title(doc)

    doc.add_heading("1. Resumen Ejecutivo y Contexto", level=1)
    doc.add_paragraph(
        "Este informe consolida dos fases de prueba aplicadas sobre la regla hasAlreadyAnswered "
        "del modulo de aprendizaje del proyecto PAE. El fragmento evaluado evita que un usuario "
        "responda dos veces la misma pregunta activa en una partida en vivo, comparando el usuario "
        "y la pregunta de cada respuesta registrada."
    )
    doc.add_paragraph(
        "El objetivo del analisis es demostrar que las pruebas de caja blanca y las pruebas "
        "basadas en experiencia no compiten: se complementan. La cobertura de condiciones aporta "
        "rigor sobre la estructura logica; el Error Guessing tensiona supuestos de entrada que la "
        "matematica de la cobertura no exige explorar."
    )
    doc.add_heading("Codigo Base Evaluado", level=2)
    add_code_block(doc, source_code)

    doc.add_heading("2. Fase 1: Analisis Formal (Cobertura de Condiciones)", level=1)
    doc.add_paragraph(
        "La Cobertura de Condiciones es una tecnica de caja blanca que busca que cada condicion "
        "booleana atomica dentro de una decision compuesta tome los valores verdadero y falso al "
        "menos una vez durante la ejecucion de las pruebas."
    )
    doc.add_paragraph(
        "En este caso, la decision compuesta es C1 && C2, donde C1 verifica si el usuario de la "
        "respuesta coincide con el usuario actual, y C2 verifica si la pregunta de la respuesta "
        "coincide con la pregunta activa."
    )
    doc.add_heading("Tabla de Verdad Estrategica", level=2)
    add_truth_table(doc)
    doc.add_heading("Resultados Empiricos", level=2)
    add_formal_results_table(doc)
    doc.add_paragraph(
        "La ejecucion de los Casos 1, 2 y 3 confirmo que el codigo supera el objetivo formal de "
        "Cobertura de Condiciones sobre la decision seleccionada: C1 fue evaluada como V/F y C2 "
        "como V/F. No se observaron fallos en esta fase cuando las entradas tenian estructura valida."
    )
    add_code_block(doc, condition_log)

    doc.add_heading("3. Fase 2: El Punto Ciego (Error Guessing)", level=1)
    doc.add_paragraph(
        "El Error Guessing es una tecnica exploratoria basada en experiencia. Su valor consiste "
        "en imaginar fallos probables por datos corruptos, valores nulos, tipos inesperados o "
        "supuestos fragiles que una tecnica formal podria no exigir."
    )
    doc.add_heading("El Caso 4 (La Ruptura)", level=2)
    doc.add_paragraph(
        "El Caso 4 inyecto una respuesta corrupta dentro del arreglo answers: answers: [null]. "
        "Este dato no altera directamente la tabla C1/C2; mas bien impide que esas condiciones "
        "se evalúen de forma segura, porque el codigo intenta acceder a answer.userId sin validar "
        "que answer sea un objeto."
    )
    doc.add_heading("Evidencia de Falla", level=2)
    doc.add_paragraph(
        "La ejecucion empirica fallo con TypeError: Cannot read properties of null (reading 'userId'). "
        "La suite reporto 3 pruebas aprobadas y 1 prueba fallida, lo que confirma el contraste: "
        "la metrica formal fue superada, pero el sistema no era robusto ante una entrada corrupta."
    )
    add_code_block(doc, error_log)

    doc.add_heading("4. Anexo Tecnico: Implementacion de la Suite de Pruebas", level=1)
    doc.add_paragraph(
        "El siguiente bloque corresponde al archivo completo de pruebas automatizadas, incluyendo "
        "los tres casos formales de Cobertura de Condiciones y el Caso 4 basado en Error Guessing."
    )
    add_code_block(doc, test_code)

    doc.add_heading("5. Conclusion: La Complementariedad", level=1)
    doc.add_paragraph(FINAL_CONCLUSION)

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("PAE - Reporte final QA cobertura vs heuristica").font.size = Pt(9)

    doc.save(PRIMARY_OUTPUT)
    doc.save(WORKSPACE_COPY)
    return PRIMARY_OUTPUT, WORKSPACE_COPY


if __name__ == "__main__":
    primary, workspace_copy = build_report()
    print(primary)
    print(workspace_copy)
